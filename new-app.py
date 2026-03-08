import streamlit as st
from openai import OpenAI
import os
from PyPDF2 import PdfReader
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
import time
import re  # 新增：正则表达式，用于提取公式

# ===================== 页面基础配置（UI优化）=====================
st.set_page_config(
    page_title="智能文档问答助手 | RAG ChatPDF",  # 页面标题
    page_icon="📄",  # 页面图标
    layout="wide",  # 宽屏布局
    initial_sidebar_state="expanded"  # 侧边栏默认展开
)

# 自定义CSS（美化界面）
st.markdown("""
<style>
/* 全局字体 */
body {font-family: "Microsoft YaHei", sans-serif;}
/* 按钮样式 */
.stButton>button {
    background-color: #4CAF50;
    color: white;
    border-radius: 8px;
    height: 3em;
    font-size: 16px;
}
.stButton>button:hover {
    background-color: #45a049;
}
/* 聊天气泡 */
.stChatMessage {
    border-radius: 12px;
    padding: 10px 15px;
    margin-bottom: 8px;
}
/* 成功/错误提示 */
.stSuccess {background-color: #f0fff4; border-radius: 8px;}
.stError {background-color: #fff5f5; border-radius: 8px;}
.stWarning {background-color: #fffbf0; border-radius: 8px;}
/* 标题样式 */
h1, h2, h3 {color: #2e4057;}
/* 侧边栏 */
.sidebar .sidebar-content {background-color: #f8f9fa;}
</style>
""", unsafe_allow_html=True)

# 核心配置：启用 KaTeX 渲染复杂 LaTeX 公式
# 引入 KaTeX 官方资源（最新版，完整功能）+ 全局样式
st.markdown("""
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css" integrity="sha384-n8MVd4RsNIU0tAv4ct0nTaAbDJwPJzDEaqSD1odI+WdtXRGWt2kTvGFasHpSy3SV" crossorigin="anonymous">
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js" integrity="sha384-XjKyOOlGwcjNTAIQHIpgOno0Hl1YQqzUOEleOLALmuqehneUG+vnGctmUb0ZY0l8" crossorigin="anonymous"></script>
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js" integrity="sha384-+VBxd3r6XgURycqtZ117nYw44OOcIax56Z4dCRWbxyPt0Koah1uHoK0o4+/RRE05" crossorigin="anonymous"></script>
<script>
    document.addEventListener("DOMContentLoaded", function() {
        renderMathInElement(document.body, {
            delimiters: [
                {left: "$$", right: "$$", display: true},  // 块级公式（居中，大字体）
                {left: "$", right: "$", display: false},   // 行内公式（嵌入文本）
                {left: "\\(", right: "\\)", display: false},
                {left: "\\[", right: "\\]", display: true}
            ],
            throwOnError: false,  // 忽略解析错误，避免公式不显示
            errorColor: "#ff0000", // 错误公式标红（便于排查）
            strict: "ignore"      // 兼容非标准LaTeX语法
        });
    });
</script>
<style>
/* 优化公式显示样式 */
.katex-display {
    font-size: 1.3em !important;  /* 放大公式（适配页面） */
    margin: 15px 0 !important;    /* 增加上下间距 */
    overflow-x: auto !important;  /* 长公式横向滚动，避免截断 */
    padding: 5px 10px !important;
}
.katex {
    font-size: 1.1em !important;  /* 行内公式放大 */
}
/* 修复Streamlit容器冲突 */
div[data-testid="stMarkdownContainer"] .katex {
    white-space: normal !important;
}
</style>
""", unsafe_allow_html=True)
# ===================== 初始化配置 =====================
# 读取秘钥
try:
    OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
    OPENAI_BASE_URL = st.secrets["OPENAI_BASE_URL"]
except KeyError:
    st.error("⚠️ 未配置API密钥！请在Streamlit Secrets中配置OPENAI_API_KEY和OPENAI_BASE_URL")
    st.stop()

# 初始化OpenAI客户端
try:
    client = OpenAI(
        api_key=OPENAI_API_KEY,
        base_url=OPENAI_BASE_URL
    )
except Exception as e:
    st.error(f"⚠️ OpenAI客户端初始化失败：{str(e)[:100]}")
    st.stop()

# 初始化会话状态
if "uploaded_files_list" not in st.session_state:
    st.session_state.uploaded_files_list = []
if "all_text" not in st.session_state:
    st.session_state.all_text = ""
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "vector_db" not in st.session_state:
    st.session_state.vector_db = None

# ===================== 核心函数 =====================
def load_file(file):
    """加载文件，带友好错误提示"""
    try:
        file_ext = os.path.splitext(file.name)[1].lower()
        text = ""
        
        if file_ext == ".pdf":
            pdf_reader = PdfReader(file)
            for idx, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text() or ""
                text += page_text + "\n\n"
            st.success(f"✅ PDF文件《{file.name}》加载完成（共{len(pdf_reader.pages)}页）")
        
        elif file_ext == ".txt":
            text = file.read().decode("utf-8")
            st.success(f"✅ TXT文件《{file.name}》加载完成（{len(text)}字符）")
        
        elif file_ext == ".docx":
            doc = Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
            st.success(f"✅ DOCX文件《{file.name}》加载完成（{len(doc.paragraphs)}段文字）")
        
        else:
            st.error(f"❌ 不支持的文件格式：{file_ext}，仅支持PDF/TXT/DOCX")
            return ""
        
        return text
    except Exception as e:
        st.error(f"❌ 文件《{file.name}》加载失败：{str(e)[:100]}")
        return ""

def init_vector_db(all_text):
    """初始化向量库，带加载提示"""
    if not all_text:
        st.warning("⚠️ 没有可处理的文本内容！")
        return None
    
    with st.spinner("🔄 正在处理文本并生成向量库...（请稍等，文本越长耗时越久）"):
        time.sleep(1)  # 加载动画更友好
        try:
            # 动态切块策略
            text_length = len(all_text)
            if text_length < 5000:
                chunk_size = 500
                chunk_overlap = 100
            elif text_length < 20000:
                chunk_size = 1000
                chunk_overlap = 200
            else:
                chunk_size = 2000
                chunk_overlap = 300
            
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", "。", "！", "？", "，", "、"]
            )
            texts = splitter.split_text(all_text)
            
            embedding = OpenAIEmbeddings(
                openai_api_key=OPENAI_API_KEY,
                openai_api_base=OPENAI_BASE_URL
            )
            db = Chroma.from_texts(texts, embedding)
            st.success(f"✅ 向量库初始化完成（分割为{len(texts)}个文本块）")
            return db
        except Exception as e:
            st.error(f"❌ 向量库初始化失败：{str(e)[:100]}")
            return None

def rag_answer(vector_db, question, chat_history, top_k=5, temp=0.1):
    """核心问答函数，新增公式自动提取+渲染"""
    if not vector_db:
        st.error("⚠️ 向量库未初始化！请先上传并加载文件")
        return ""
    
    if not question or question.strip() == "":
        st.warning("⚠️ 请输入有效的问题！")
        return ""
    
    with st.spinner("🔍 正在检索相关内容并生成回答..."):
        try:
            # 1. 检索相关文档
            relevant_docs = vector_db.similarity_search(query=question, k=top_k)
            if not relevant_docs:
                st.warning("⚠️ 未检索到相关内容，请调整问题或增加检索条数")
                return "未在文档中找到相关信息"
            
            # 2. 拼接上下文（保留原始公式格式）
            context = ""
            for idx, doc in enumerate(relevant_docs[:3]):
                context += f"相关内容{idx+1}：\n{doc.page_content}\n\n"
            
            # 3. 调用大模型（要求返回公式时用LaTeX格式）
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                temperature=temp,
                messages=[
                    {"role": "system", "content": """你是一个专业的文档问答助手，严格遵守以下规则：
1. 仅使用提供的「相关内容」回答问题，绝对不编造任何信息；
2. 如果回答中包含公式，必须用标准LaTeX格式书写（行内公式用$...$，块级公式用$$...$$）；
3. 公式中的上标用^、下标用_，分式用\\frac，偏导数用\\partial，确保LaTeX语法正确；
4. 结合历史对话理解上下文，回答要简洁、准确、有针对性；
5. 回答使用中文，格式清晰，重点内容可加粗。"""},
                    {"role": "user", "content": f"""相关内容（仅基于此回答）：
{context}

历史对话（仅用于理解上下文）：
{chat_history}

当前问题：
{question}"""}
                ]
            )
            answer = response.choices[0].message.content.strip()
            
            # 4. 返回原始回答（后续在显示环节拆分公式+渲染）
            return answer if answer else "未在文档中找到相关信息"
        
        except Exception as e:
            error_msg = str(e)[:200]
            st.error(f"❌ 回答生成失败：{error_msg}")
            st.info("建议检查：1.API Key是否有效 2.Base URL是否正确 3.网络是否正常")
            return ""

# ========== 新增：公式拆分+渲染函数 ==========
def clean_text_for_rendering(text):
    """
    清理文本中多余的换行符和空白，避免公式被错误拆分：
    - 合并连续的换行符（保留段落间的空行）
    - 清理公式前后的换行，让公式和文本自然衔接
    """
    if not text:
        return text
    
    # 1. 合并连续的换行符（保留段落间的空行）
    # 先把所有换行替换为 \n，再合并连续的 \n 为单个 \n
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 合并连续的换行（保留段落间的空行）
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # 2. 清理公式前后的换行（关键！避免公式被拆成多行）
    # 匹配块级公式 $$...$$ 前后的换行，清理掉
    text = re.sub(r'\n+(\$\$.*?\$\$)\n+', r' \1 ', text, flags=re.DOTALL)
    # 匹配行内公式 $...$ 前后的换行，清理掉
    text = re.sub(r'\n+(\$.*?\$)\n+', r' \1 ', text, flags=re.DOTALL)
    
    # 3. 清理多余的空格和制表符
    text = re.sub(r' +', ' ', text)  # 合并连续空格
    text = re.sub(r'\t+', ' ', text) # 替换制表符为空格
    
    return text.strip()

def render_text_with_formulas(text):
    """
    优化后的公式渲染函数：
    - 先清理文本中的多余换行
    - 行内公式用 markdown 包裹，保持和文本同行
    - 块级公式居中显示，避免不必要的换行
    """
    if not text:
        return
    
    # 关键：先清理文本中的多余换行
    cleaned_text = clean_text_for_rendering(text)
    
    # 正则表达式：优先匹配块级公式，避免行内公式覆盖
    pattern = r'(\$\$.*?\$\$)|(\$.*?\$)'
    parts = re.split(pattern, cleaned_text, flags=re.DOTALL)
    
    # 用于拼接行内公式和前后文本，避免不必要的换行
    current_text = ""
    
    for part in parts:
        if part is None or part == "":
            continue
        
        # 块级公式（$$...$$）
        if part.startswith("$$") and part.endswith("$$"):
            # 先输出当前累积的文本
            if current_text.strip():
                st.markdown(current_text, unsafe_allow_html=True)
                current_text = ""
            # 渲染块级公式
            formula = part.strip("$$").replace("\n", "").strip()
            st.latex(formula)
        # 行内公式（$...$）
        elif part.startswith("$") and part.endswith("$"):
            # 行内公式直接拼接到 current_text 中，保持和文本同行
            formula = part.strip("$").strip()
            current_text += f" ${formula}$ "
        # 普通文本
        else:
            current_text += part
    
    # 输出最后一段累积的文本
    if current_text.strip():
        st.markdown(current_text, unsafe_allow_html=True)

# ===================== 界面布局 =====================
# 主标题
st.title("📄 智能文档问答助手")
st.markdown("---")

# 侧边栏配置
with st.sidebar:
    st.header("⚙️ 配置中心")
    st.markdown("### 检索参数")
    top_k = st.slider("检索条数（K值）", min_value=3, max_value=10, value=5, 
                     help="值越大召回率越高，但提示词越长（推荐5-7）")
    temp = st.slider("回答随机性（温度）", min_value=0.0, max_value=0.5, value=0.1, 
                    help="值越低回答越准确，越高越灵活（推荐0.1）")
    
    st.markdown("---")
    st.markdown("### 📖 使用说明")
    st.markdown("""
    1. 上传PDF/TXT/DOCX格式文件（可多选）
    2. 点击「添加文件」完成加载
    3. 输入问题，支持多轮上下文对话
    4. 可调整检索参数优化回答效果
    
    ✨ 特点：
    - 支持多格式文件
    - 多轮上下文记忆
    - 低幻觉，高准确率
    - 中文语义优化
    """)
    st.markdown("---")
    st.markdown("#### 🚀 技术栈")
    st.markdown("Python / Streamlit / LangChain / OpenAI / Chroma")

# 主内容区
col1, col2 = st.columns([2, 1])

with col1:
    # 文件上传区域
    st.subheader("📤 文件上传")
    uploaded_file = st.file_uploader(
        "选择文件（支持PDF/TXT/DOCX）",
        type=["pdf", "txt", "docx"],
        label_visibility="collapsed"
    )
    
    # 按钮行
    btn_col1, btn_col2 = st.columns(2)
    with btn_col1:
        if st.button("➕ 添加文件", use_container_width=True):
            if uploaded_file:
                # 检查重复
                file_names = [f.name for f in st.session_state.uploaded_files_list]
                if uploaded_file.name in file_names:
                    st.warning(f"⚠️ 文件《{uploaded_file.name}》已上传，无需重复添加！")
                else:
                    # 加载文件
                    file_text = load_file(uploaded_file)
                    if file_text:
                        st.session_state.uploaded_files_list.append(uploaded_file)
                        st.session_state.all_text += file_text + "\n\n"
                        # 重新初始化向量库
                        st.session_state.vector_db = init_vector_db(st.session_state.all_text)
            else:
                st.warning("⚠️ 请先选择要上传的文件！")
    
    with btn_col2:
        if st.button("🗑️ 清空所有文件", use_container_width=True, type="secondary"):
            st.session_state.uploaded_files_list = []
            st.session_state.all_text = ""
            st.session_state.vector_db = None
            st.success("✅ 已清空所有文件和向量库！")
    
    st.markdown("---")

with col2:
    # 已上传文件列表
    st.subheader("📋 已上传文件")
    if st.session_state.uploaded_files_list:
        for idx, f in enumerate(st.session_state.uploaded_files_list):
            st.markdown(f"{idx+1}. 📄 {f.name}")
        st.markdown(f"**总计：{len(st.session_state.uploaded_files_list)} 个文件**")
        
        # 清空对话历史按钮
        if st.button("🧹 清空对话历史", type="secondary"):
            st.session_state.chat_history = []
            st.success("✅ 已清空对话历史！")
            st.balloons()
    else:
        st.info("暂无上传文件，请先添加文件后提问")

# 对话区域
st.markdown("---")
st.subheader("💬 对话区域")

# 显示对话历史
if st.session_state.chat_history:
    for q, a in st.session_state.chat_history:
        st.chat_message("user").write(f"**你**：{q}")
        st.chat_message("assistant").write(f"**助手**：{a}")
        render_text_with_formulas(a)

        
# 提问区域
question = st.text_input("请输入你的问题（支持上下文对话）：", placeholder="比如：文档中提到的核心功能是什么？", label_visibility="collapsed")
if st.button("🚀 获取答案", use_container_width=True):
    if st.session_state.uploaded_files_list and st.session_state.vector_db:
        answer = rag_answer(
            vector_db=st.session_state.vector_db,
            question=question,
            chat_history=st.session_state.chat_history,
            top_k=top_k,
            temp=temp
        )
        if answer:
            # 显示回答
            # 关键修改：用自定义函数渲染（自动拆分公式）
            st.chat_message("assistant").write("**助手**：")
            render_text_with_formulas(answer)
            # 保存到历史
            st.session_state.chat_history.append((question, answer))
    else:
        st.error("⚠️ 请先上传并加载文件后再提问！")

# 页脚
st.markdown("---")
st.markdown("© 2026 智能文档问答助手 | 基于 RAG 架构构建 | [GitHub 仓库](https://github.com/GrandYK/chatPDF)")